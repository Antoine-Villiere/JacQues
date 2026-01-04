from __future__ import annotations

from pathlib import Path
from typing import List
from datetime import datetime, timezone
import json
import re
import time
from urllib.parse import quote, urlencode
import sqlite3
import threading
import warnings
import secrets
import shutil
import subprocess
import sys

warnings.filterwarnings(
    "ignore",
    message="Valid config keys have changed in V2:",
    category=UserWarning,
)

from dotenv import load_dotenv
from flask import abort, send_from_directory, request, jsonify
import requests
import dash
from dash import dcc, html, dash_table, Input, Output, State, no_update, ALL, MATCH

from jacques import db
from jacques.config import (
    BASE_DIR,
    DATA_DIR,
    EXPORTS_DIR,
    IMAGES_DIR,
    UPLOADS_DIR,
    Settings,
    ensure_dirs,
)
from jacques.services import assistant, doc_ingest, file_ops, pdf_tools, rag, vision, scheduler as task_scheduler
from jacques.utils import decode_upload, safe_filename


load_dotenv(BASE_DIR / ".env", override=True)
ensure_dirs()

db.init_db()
settings = Settings()
STREAMING_STATUS: dict[int, bool] = {}
STREAMING_LOCK = threading.Lock()
TOOL_STATUS: dict[int, dict[str, str]] = {}
TOOL_LOCK = threading.Lock()
CANCEL_STATUS: dict[int, bool] = {}
CANCEL_LOCK = threading.Lock()
MAILTO_SEEN: dict[int, int] = {}
MAILTO_LOCK = threading.Lock()
MAILTO_RECENT_SECONDS = 300
CALENDAR_SEEN: dict[int, int] = {}
CALENDAR_LOCK = threading.Lock()
CALENDAR_RECENT_SECONDS = 300
SCHEDULER_STARTED = False
SCHEDULER_LOCK = threading.Lock()
INITIAL_SYSTEM_PROMPT = db.get_setting("system_prompt") or assistant.SYSTEM_PROMPT
INITIAL_GLOBAL_MEMORY = db.get_setting("global_memory") or ""
INITIAL_CUSTOM_INSTRUCTIONS = db.get_setting("custom_instructions") or ""
INITIAL_USER_NICKNAME = db.get_setting("user_nickname") or ""
INITIAL_USER_OCCUPATION = db.get_setting("user_occupation") or ""
INITIAL_USER_ABOUT = db.get_setting("user_about") or ""
INITIAL_DEFAULT_CALENDAR = db.get_setting("default_calendar") or ""
INITIAL_MAIL_ACCOUNT = db.get_setting("default_mail_account") or ""
INITIAL_MAILBOX = db.get_setting("default_mailbox") or ""
INITIAL_ONLYOFFICE_URL = db.get_setting("onlyoffice_url") or settings.onlyoffice_url or ""
INITIAL_ONLYOFFICE_JWT = db.get_setting("onlyoffice_jwt") or settings.onlyoffice_jwt or ""
INITIAL_APP_BASE_URL = db.get_setting("app_base_url") or settings.app_base_url or ""
INITIAL_APP_TIMEZONE = db.get_setting("app_timezone") or settings.app_timezone or "UTC"
INITIAL_LLM_STREAMING = db.get_setting("llm_streaming")


def _setting_bool(value: str | None, default: bool) -> bool:
    if value is None:
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def _setting_enabled(key: str, default: bool = True) -> bool:
    return _setting_bool(db.get_setting(key), default)


TOOLS_WEB_ENABLED = _setting_bool(db.get_setting("tools_web_enabled"), True)
TOOLS_MAIL_ENABLED = _setting_bool(db.get_setting("tools_mail_enabled"), True)
TOOLS_CALENDAR_ENABLED = _setting_bool(db.get_setting("tools_calendar_enabled"), True)
TOOLS_CODE_ENABLED = _setting_bool(db.get_setting("tools_code_enabled"), True)
LLM_STREAMING_ENABLED = _setting_bool(INITIAL_LLM_STREAMING, settings.llm_streaming)
if INITIAL_ONLYOFFICE_URL:
    settings.onlyoffice_url = INITIAL_ONLYOFFICE_URL
if INITIAL_ONLYOFFICE_JWT:
    settings.onlyoffice_jwt = INITIAL_ONLYOFFICE_JWT
if INITIAL_APP_BASE_URL:
    settings.app_base_url = INITIAL_APP_BASE_URL
if INITIAL_APP_TIMEZONE:
    settings.app_timezone = INITIAL_APP_TIMEZONE
settings.llm_streaming = LLM_STREAMING_ENABLED


def _set_streaming(conversation_id: int, value: bool) -> None:
    with STREAMING_LOCK:
        STREAMING_STATUS[conversation_id] = value


def _is_streaming(conversation_id: int) -> bool:
    with STREAMING_LOCK:
        return STREAMING_STATUS.get(conversation_id, False)


def _set_cancelled(conversation_id: int, value: bool) -> None:
    with CANCEL_LOCK:
        CANCEL_STATUS[conversation_id] = value


def _is_cancelled(conversation_id: int) -> bool:
    with CANCEL_LOCK:
        return CANCEL_STATUS.get(conversation_id, False)


def _set_tool_status(conversation_id: int, name: str | None, stage: str) -> None:
    with TOOL_LOCK:
        if not name:
            TOOL_STATUS.pop(conversation_id, None)
            return
        TOOL_STATUS[conversation_id] = {"name": name, "stage": stage}


def _get_tool_status(conversation_id: int) -> dict[str, str] | None:
    with TOOL_LOCK:
        return TOOL_STATUS.get(conversation_id)


def _extract_tool_call_args(content: str) -> dict | None:
    if not content:
        return None
    match = re.search(r"```json\\n(.*?)\\n```", content, re.DOTALL)
    if not match:
        return None
    raw = match.group(1).strip()
    if not raw:
        return None
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return None


def _extract_tool_result_text(content: str) -> str:
    if not content:
        return ""
    if "```text" in content:
        try:
            return content.split("```text", 1)[1].split("```", 1)[0].strip()
        except Exception:
            return ""
    parts = content.split("\n", 1)
    return parts[1].strip() if len(parts) > 1 else content.strip()


def _normalize_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    if not text:
        return []
    parts = [item.strip() for item in text.replace(";", ",").split(",")]
    return [part for part in parts if part]


def _build_mailto(args: dict | None) -> str | None:
    if not isinstance(args, dict):
        return None
    to_list = _normalize_list(args.get("to"))
    cc_list = _normalize_list(args.get("cc"))
    bcc_list = _normalize_list(args.get("bcc"))
    subject = str(args.get("subject") or "").strip()
    body = str(args.get("body") or "").strip()
    if not to_list and not subject and not body and not cc_list and not bcc_list:
        return None
    mailto = "mailto:" + ",".join(to_list)
    params: dict[str, str] = {}
    if subject:
        params["subject"] = subject
    if body:
        params["body"] = body
    if cc_list:
        params["cc"] = ",".join(cc_list)
    if bcc_list:
        params["bcc"] = ",".join(bcc_list)
    if params:
        mailto = f"{mailto}?{urlencode(params, quote_via=quote)}"
    return mailto


def _should_open_mailto(conversation_id: int, message_id: int) -> bool:
    with MAILTO_LOCK:
        last_seen = MAILTO_SEEN.get(conversation_id)
        if last_seen == message_id:
            return False
        MAILTO_SEEN[conversation_id] = message_id
        return True


def _should_open_calendar(conversation_id: int, message_id: int) -> bool:
    with CALENDAR_LOCK:
        last_seen = CALENDAR_SEEN.get(conversation_id)
        if last_seen == message_id:
            return False
        CALENDAR_SEEN[conversation_id] = message_id
        return True


def _is_recent_message(created_at: str | None, max_age_seconds: int) -> bool:
    if not created_at:
        return False
    try:
        created = datetime.fromisoformat(created_at)
    except ValueError:
        return False
    if created.tzinfo is None:
        created = created.replace(tzinfo=timezone.utc)
    delta = datetime.now(timezone.utc) - created.astimezone(timezone.utc)
    return delta.total_seconds() <= max_age_seconds


def _extract_ics_link(content: str) -> str | None:
    if not content:
        return None
    match = re.search(r"(\/files\/generated\/[\\w\\-\\.]+\\.ics)", content)
    if not match:
        return None
    return match.group(1)


def _ensure_scheduler_started() -> None:
    global SCHEDULER_STARTED
    with SCHEDULER_LOCK:
        if SCHEDULER_STARTED:
            return
        task_scheduler.start(settings)
        SCHEDULER_STARTED = True


def _ensure_default_conversation() -> int:
    conversations = db.list_conversations()
    if conversations:
        return int(conversations[0]["id"])
    return db.create_conversation("Conversation 1")


def _conversation_options() -> list[dict[str, str]]:
    return [
        {
            "label": html.Div(
                [
                    html.Span(row["title"], className="convo-title"),
                    html.Button(
                        "×",
                        id={"type": "delete-convo", "index": row["id"]},
                        className="icon-btn convo-delete",
                        n_clicks=0,
                    ),
                ],
                className="convo-item",
            ),
            "value": str(row["id"]),
        }
        for row in db.list_conversations()
    ]


def _purge_conversation(conversation_id: int) -> None:
    if _is_streaming(conversation_id):
        _set_streaming(conversation_id, False)
    _set_tool_status(conversation_id, None, "idle")
    docs = db.list_documents(conversation_id)
    images = db.list_images(conversation_id)
    tasks = db.list_scheduled_tasks(conversation_id)
    db.delete_conversation(conversation_id)
    rag.delete_index(conversation_id)
    for task in tasks:
        task_scheduler.remove_task(int(task["id"]))
    for doc in docs:
        _maybe_delete_file(doc["path"], "documents")
    for img in images:
        _maybe_delete_file(img["path"], "images")


def _export_all_data() -> Path:
    export_payload: dict[str, object] = {
        "exported_at": datetime.now(timezone.utc).isoformat(),
    }
    with db.get_connection() as conn:
        for table in [
            "conversations",
            "messages",
            "documents",
            "images",
            "pdf_highlights",
            "scheduled_tasks",
            "app_settings",
        ]:
            cursor = conn.execute(f"SELECT * FROM {table}")
            export_payload[table] = [dict(row) for row in cursor.fetchall()]
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"jacques_export_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
    path = EXPORTS_DIR / filename
    path.write_text(json.dumps(export_payload, ensure_ascii=True, indent=2))
    return path


def _message_class(role: str) -> str:
    if role in {"user", "assistant", "tool"}:
        return f"message message-{role}"
    return "message message-tool"


def _tool_message_is_error(content: str) -> bool:
    lowered = content.lower()
    markers = [
        "failed",
        "error",
        "invalid tool arguments",
        "unknown tool",
        "exception",
        "traceback",
    ]
    return any(marker in lowered for marker in markers)


def _render_messages(rows: List[sqlite3.Row]):
    if not rows:
        return [html.Div("No messages yet.", className="status")]
    rendered = []
    for row in rows:
        role = row["role"]
        if role == "tool" and not _tool_message_is_error(row["content"] or ""):
            continue
        rendered.append(
            html.Div(
                className=_message_class(role),
                children=[
                    html.Div(role.upper(), className="message-role"),
                    dcc.Markdown(
                        row["content"],
                        className="message-content",
                        link_target="_blank",
                    ),
                ],
            )
        )
    return rendered


def _doc_label(doc_type: str) -> tuple[str, str]:
    mapping = {
        ".pdf": ("PDF", "pdf"),
        ".docx": ("DOC", "doc"),
        ".xlsx": ("XLS", "xls"),
        ".xls": ("XLS", "xls"),
        ".xlsm": ("XLS", "xls"),
        ".csv": ("CSV", "csv"),
    }
    return mapping.get(doc_type, ("DOC", "file"))


def _asset_pill(label: str, name: str, kind: str) -> html.Div:
    return html.Div(
        className=f"asset-pill asset-pill--{kind}",
        children=[
            html.Span(label, className="asset-kind"),
            html.Span(name, className="asset-name"),
        ],
    )


def _parse_task_payload(raw: str | None) -> dict:
    if not raw:
        return {}
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {}


def _task_payload_summary(task: sqlite3.Row) -> str:
    payload = _parse_task_payload(task["payload"])
    if task["task_type"] == "reminder":
        return str(payload.get("message") or "").strip()
    if task["task_type"] == "web_digest":
        return str(payload.get("query") or "").strip()
    return ""


def _get_docs_for_conversation(convo_id: str | None) -> list[dict[str, str | int]]:
    if not convo_id:
        return []
    try:
        conversation_id = int(convo_id)
    except ValueError:
        return []
    docs = db.list_documents(conversation_id)
    return [
        {"id": int(doc["id"]), "name": doc["name"], "doc_type": doc["doc_type"]}
        for doc in docs
    ]


def _apply_doc_mention(text: str, doc_name: str) -> str:
    mention = f'@"{doc_name}" '
    if not text:
        return mention
    trimmed = text.rstrip()
    match = re.search(r"@[^\\s@]*$", trimmed)
    if match:
        return f"{trimmed[: match.start()]}{mention}"
    spacer = "" if trimmed.endswith(" ") or not trimmed else " "
    return f"{trimmed}{spacer}{mention}"


def _file_item(label: str, name: str, detail: str, item_id: dict) -> html.Button:
    return html.Button(
        children=[
            html.Div(label, className="file-item-kind"),
            html.Div(
                [
                    html.Div(name, className="file-item-name"),
                    html.Div(detail, className="file-item-meta"),
                ],
                className="file-item-text",
            ),
        ],
        id=item_id,
        n_clicks=0,
        className="file-item",
    )


def _stream_response(
    conversation_id: int,
    user_message: str,
    use_rag: bool,
    use_web: bool,
    assistant_message_id: int,
) -> None:
    streamed = False

    def on_token(token: str) -> None:
        nonlocal streamed
        streamed = True
        db.append_message_content(assistant_message_id, token)

    def on_tool_event(name: str, stage: str) -> None:
        _set_tool_status(conversation_id, name, stage)

    reply = ""
    try:
        reply = assistant.respond_streaming(
            conversation_id,
            user_message,
            settings,
            use_rag=use_rag,
            use_web=use_web,
            on_token=on_token,
            on_tool_event=on_tool_event,
            should_cancel=lambda: _is_cancelled(conversation_id),
        )
    except Exception as exc:
        reply = f"LLM error: {exc}"
    finally:
        if reply:
            db.update_message_content(assistant_message_id, reply)
        elif not streamed:
            db.update_message_content(assistant_message_id, "No response generated.")
        assistant.maybe_update_conversation_title(
            conversation_id, settings, force_first=True
        )
        _set_streaming(conversation_id, False)
        _set_cancelled(conversation_id, False)


def _latest_tool_name(conversation_id: int) -> str | None:
    row = db.get_latest_tool_message(conversation_id)
    if not row:
        return None
    content = row["content"] or ""
    match = re.search(r"Tool (?:call|result): \\*\\*(.+?)\\*\\*", content)
    if match:
        return match.group(1).strip()
    return None


def _maybe_delete_file(path: str, table: str) -> None:
    if not path:
        return
    if db.is_path_used(table, path):
        return
    file_path = Path(path)
    if file_path.exists():
        file_path.unlink()


def _data_url(path: str) -> str:
    if not path:
        return ""
    resolved = Path(path).resolve()
    try:
        relative = resolved.relative_to(DATA_DIR)
    except ValueError:
        return ""
    return f"/files/{quote(relative.as_posix(), safe='/')}"


def _absolute_file_url(path: str, base_url: str | None = None) -> str:
    relative = _data_url(path)
    if not relative:
        return ""
    base = (base_url or settings.app_base_url or "").rstrip("/")
    if not base:
        return relative
    return f"{base}{relative}"


def _effective_onlyoffice_base_url() -> str:
    base = (settings.app_base_url or "").strip() or "http://127.0.0.1:8050"
    onlyoffice = (settings.onlyoffice_url or "").strip()
    if onlyoffice and ("localhost" in onlyoffice or "127.0.0.1" in onlyoffice):
        if "localhost" in base or "127.0.0.1" in base:
            base = _default_docker_base_url()
    return base.rstrip("/")


def _run_command(args: list[str]) -> tuple[bool, str]:
    try:
        result = subprocess.run(
            args,
            check=False,
            capture_output=True,
            text=True,
        )
    except Exception as exc:
        return False, str(exc)
    output = (result.stdout or "").strip()
    err = (result.stderr or "").strip()
    if result.returncode != 0:
        return False, err or output or "Command failed."
    return True, output or "OK"


def _start_onlyoffice_container(jwt_secret: str) -> tuple[bool, str]:
    if not shutil.which("docker"):
        return False, "Docker is not installed or not in PATH."
    container_name = "jacques-onlyoffice"
    _run_command(["docker", "rm", "-f", container_name])
    cmd = [
        "docker",
        "run",
        "-d",
        "--name",
        container_name,
        "-p",
        "8080:80",
        "-e",
        "JWT_ENABLED=true",
        "-e",
        f"JWT_SECRET={jwt_secret}",
        "-e",
        "JWT_HEADER=Authorization",
        "onlyoffice/documentserver",
    ]
    ok, output = _run_command(cmd)
    if not ok:
        return False, output
    return True, "OnlyOffice started on http://localhost:8080"


def _start_docker_desktop() -> tuple[bool, str]:
    if shutil.which("docker"):
        ok, _ = _run_command(["docker", "info"])
        if ok:
            return True, "Docker is already running."
    if shutil.which("open"):
        ok, message = _run_command(["open", "-a", "Docker"])
        if ok:
            return True, "Starting Docker Desktop..."
        return False, message
    return False, "Cannot launch Docker Desktop on this OS."


def _default_docker_base_url() -> str:
    if sys.platform.startswith("darwin") or sys.platform.startswith("win"):
        return "http://host.docker.internal:8050"
    return "http://172.17.0.1:8050"


def _load_word_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word editing") from exc
    doc = Document(str(path))
    lines = [para.text for para in doc.paragraphs]
    return "\n".join(lines).strip()


def _format_word_preview(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word editing") from exc
    doc = Document(str(path))
    lines = []
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            lines.append(f"{idx}. {text}")
    return "\n".join(lines).strip() or "No text content found."


def _format_word_preview_html(path: Path) -> str:
    try:
        import mammoth
    except ImportError:
        return ""
    with path.open("rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
    html_body = (result.value or "").strip()
    if not html_body:
        return ""
    return f"""<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <style>
      body {{
        font-family: "Times New Roman", serif;
        color: #111;
        background: #fff;
        padding: 20px 24px;
        line-height: 1.5;
      }}
      h1, h2, h3, h4, h5, h6 {{
        font-family: "Times New Roman", serif;
        margin: 0 0 0.6em;
      }}
      p {{
        margin: 0 0 0.85em;
      }}
      table {{
        border-collapse: collapse;
        width: 100%;
        margin: 0 0 1em;
      }}
      td, th {{
        border: 1px solid #ddd;
        padding: 6px 8px;
      }}
    </style>
  </head>
  <body>
    {html_body}
  </body>
</html>"""


def _render_word_preview(path: Path):
    html_doc = _format_word_preview_html(path)
    if html_doc:
        return html.Iframe(srcDoc=html_doc, className="docx-frame")
    preview = _format_word_preview(path)
    return html.Pre(preview, className="file-preview")


def _load_table_data(path: Path, sheet_name: str | None = None):
    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError("pandas is required for table editing") from exc
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
        return df, []
    if sheet_name:
        df = pd.read_excel(path, sheet_name=sheet_name)
    else:
        df = pd.read_excel(path, sheet_name=0)
    sheets = []
    try:
        xls = pd.ExcelFile(path)
        sheets = list(xls.sheet_names)
    except Exception:
        sheets = []
    return df, sheets


def _coerce_cell_value(value):
    if value is None:
        return None
    if isinstance(value, str):
        stripped = value.strip()
        return stripped if stripped else None
    return value


default_convo_id = _ensure_default_conversation()

app = dash.Dash(__name__, suppress_callback_exceptions=True, update_title="")
app.title = "Jacques Assistant"

ALLOWED_FILE_DIRS = {"generated", "images", "uploads", "exports"}


@app.server.before_request
def _start_scheduler_once():
    _ensure_scheduler_started()


@app.server.route("/files/<path:asset_path>")
def serve_files(asset_path: str):
    safe_path = Path(asset_path)
    if safe_path.is_absolute() or ".." in safe_path.parts:
        abort(404)
    if not safe_path.parts or safe_path.parts[0] not in ALLOWED_FILE_DIRS:
        abort(404)
    full_path = (DATA_DIR / safe_path).resolve()
    if not full_path.exists() or DATA_DIR not in full_path.parents:
        abort(404)
    return send_from_directory(DATA_DIR, safe_path.as_posix())


@app.server.route("/pdf_highlight", methods=["POST"])
def handle_pdf_highlight():
    payload = request.get_json(silent=True) or {}
    doc_id = int(payload.get("doc_id") or 0)
    text = str(payload.get("text") or "").strip()
    page_number = payload.get("page")
    page_value = None
    if page_number is not None and str(page_number).strip():
        try:
            page_value = int(page_number)
        except ValueError:
            page_value = None

    if not doc_id or not text:
        return jsonify({"success": False, "message": "Missing highlight data."}), 400
    doc = db.get_document_by_id(doc_id)
    if not doc:
        return jsonify({"success": False, "message": "Document not found."}), 404
    path = Path(doc["path"]).resolve()
    if path.suffix.lower() != ".pdf":
        return jsonify({"success": False, "message": "Not a PDF file."}), 400
    if DATA_DIR not in path.parents:
        return jsonify({"success": False, "message": "Invalid file path."}), 400

    try:
        count = pdf_tools.highlight_text(path, text, page_value)
    except Exception as exc:
        return jsonify({"success": False, "message": f"Highlight failed: {exc}"}), 500

    if count == 0:
        return jsonify(
            {"success": False, "message": "No matches found to highlight."}
        ), 200

    db.add_pdf_highlight(doc_id, page_value, text)
    try:
        rag.build_index(int(doc["conversation_id"]))
    except Exception:
        pass

    return jsonify(
        {"success": True, "message": f"Highlighted {count} match(es)."}
    )


@app.server.route("/onlyoffice/<int:doc_id>")
def onlyoffice_editor(doc_id: int):
    if not settings.onlyoffice_url:
        abort(404)
    doc = db.get_document_by_id(doc_id)
    if not doc:
        abort(404)

    path = Path(doc["path"])
    if not path.exists():
        abort(404)

    base_url = _effective_onlyoffice_base_url()
    file_url = _absolute_file_url(doc["path"], base_url)
    if not file_url:
        abort(404)

    file_type = path.suffix.lstrip(".").lower()
    key_seed = f"{doc_id}-{int(path.stat().st_mtime)}-{path.name}"
    key = str(abs(hash(key_seed)))
    callback_url = f"{base_url}/onlyoffice_callback?doc_id={doc_id}"

    config = {
        "document": {
            "fileType": file_type,
            "key": key,
            "title": doc["name"],
            "url": file_url,
        },
        "editorConfig": {
            "mode": "edit",
            "callbackUrl": callback_url,
        },
    }

    token_script = ""
    if settings.onlyoffice_jwt:
        try:
            import jwt
        except ImportError:
            return (
                "PyJWT is required for ONLYOFFICE_JWT. Install PyJWT.",
                500,
            )
        token = jwt.encode(config, settings.onlyoffice_jwt, algorithm="HS256")
        if isinstance(token, bytes):
            token = token.decode("utf-8")
        config["token"] = token
        token_script = f'var token = "{token}";'

    docserver = settings.onlyoffice_url.rstrip("/")
    html = f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <title>{doc["name"]}</title>
    <style>
      html, body {{
        margin: 0;
        padding: 0;
        height: 100%;
        background: #0f141b;
      }}
      #editor {{
        height: 100%;
        width: 100%;
      }}
      #onlyoffice-status {{
        position: absolute;
        inset: 16px 16px auto 16px;
        background: rgba(15, 23, 42, 0.92);
        color: #e2e8f0;
        padding: 12px 14px;
        border-radius: 12px;
        font-family: "Spline Sans Mono", monospace;
        font-size: 12px;
        max-width: 520px;
        z-index: 2;
      }}
      #onlyoffice-status a {{
        color: #7dd3fc;
      }}
    </style>
  </head>
  <body>
    <div id="editor"></div>
    <div id="onlyoffice-status">Loading OnlyOffice...</div>
    <script src="{docserver}/web-apps/apps/api/documents/api.js" onerror="window.onlyofficeLoadError = true;"></script>
    <script>
      {token_script}
      var config = {json.dumps(config)};
      var statusEl = document.getElementById("onlyoffice-status");
      function setStatus(message) {{
        statusEl.innerHTML = message;
      }}
      function tryInit() {{
        if (window.onlyofficeLoadError || !window.DocsAPI) {{
          setStatus("OnlyOffice API not loaded. Check ONLYOFFICE_URL and Docker status.");
          return;
        }}
        try {{
          window.docEditor = new DocsAPI.DocEditor("editor", config);
          setStatus("");
        }} catch (err) {{
          setStatus("OnlyOffice init failed: " + err);
        }}
      }}
      setTimeout(function () {{
        if (!window.docEditor) {{
          setStatus(
            "OnlyOffice not responding.<br/>" +
            "File URL: <a href='{file_url}' target='_blank' rel='noreferrer'>{file_url}</a><br/>" +
            "Callback URL: {callback_url}"
          );
        }}
      }}, 4000);
      tryInit();
    </script>
  </body>
</html>"""
    return html, 200, {"Content-Type": "text/html"}


@app.server.route("/onlyoffice_callback", methods=["POST"])
def onlyoffice_callback():
    payload = request.get_json(silent=True) or {}
    status = payload.get("status")
    download_url = payload.get("url")
    token = payload.get("token") or payload.get("jwt")
    doc_id = request.args.get("doc_id")
    if not doc_id:
        return jsonify({"error": 0})
    try:
        doc_id_int = int(doc_id)
    except ValueError:
        return jsonify({"error": 0})

    if status in {2, 6}:
        if not download_url:
            return jsonify({"error": 1, "message": "Missing download URL."})
        doc = db.get_document_by_id(doc_id_int)
        if doc:
            path = Path(doc["path"])
            try:
                headers = {}
                if token:
                    headers["Authorization"] = f"Bearer {token}"
                response = requests.get(
                    download_url,
                    timeout=settings.web_timeout,
                    headers=headers,
                )
                response.raise_for_status()
                path.write_bytes(response.content)
                text = doc_ingest.extract_text(path)
                db.update_document_text(doc_id_int, text)
                rag.build_index(int(doc["conversation_id"]))
            except Exception as exc:
                return jsonify({"error": 1, "message": str(exc)})
    return jsonify({"error": 0})

app.layout = html.Div(
    id="theme-root",
    className="app-shell",
    **{"data-theme": "light"},
    children=[
        dcc.Store(id="chat-refresh", data=0),
        dcc.Store(id="docs-refresh", data=0),
        dcc.Store(id="images-refresh", data=0),
        dcc.Store(
            id="stream-state",
            data={"convo_id": str(default_convo_id), "is_streaming": False},
        ),
        dcc.Store(id="settings-open", data=False),
        dcc.Store(id="active-file", data=None),
        dcc.Store(id="scroll-trigger", data=0),
        dcc.Store(id="mailto-data", data={}),
        dcc.Store(id="mailto-state", data={}),
        dcc.Store(id="calendar-data", data={}),
        dcc.Store(id="calendar-state", data={}),
        dcc.Store(id="archived-refresh", data=0),
        dcc.Store(id="tasks-refresh", data=0),
        dcc.Store(id="docs-list", data=[]),
        dcc.Interval(id="stream-interval", interval=120, n_intervals=0),
        html.Div(
            className="sidebar",
            children=[
                html.Div("Jacques", className="brand"),
                html.Div(
                    className="section",
                    children=[
                        html.Div(
                            className="section-header",
                            children=[
                                html.H3("Conversations"),
                                html.Div(
                                    className="section-actions",
                                    children=[
                                        html.Button(
                                            "+",
                                            id="new-convo-btn",
                                            className="icon-btn icon-btn--add",
                                            **{
                                                "data-tooltip": "New conversation."
                                            },
                                        ),
                                        html.Button(
                                            "SYS",
                                            id="open-settings-btn",
                                            className="icon-btn",
                                            **{
                                                "data-tooltip": "System prompt & global memory."
                                            },
                                        ),
                                    ],
                                ),
                            ],
                        ),
                        dcc.RadioItems(
                            id="convo-dropdown",
                            options=_conversation_options(),
                            value=str(default_convo_id),
                            className="convo-list",
                        ),
                        html.Div(id="convo-status", className="status"),
                    ],
                ),
            ],
        ),
        html.Div(
            className="main",
            children=[
                html.Div(
                    className="chat-header",
                    children=[
                        html.Div(
                            className="chat-title-block",
                            children=[
                                html.Div(
                                    className="chat-title-row",
                                    children=[
                                        html.Div(
                                            id="chat-title", className="chat-title"
                                        ),
                                        html.Button(
                                            "✎",
                                            id="edit-title-btn",
                                            className="icon-btn icon-btn--edit",
                                            **{
                                                "data-tooltip": "Rename conversation."
                                            },
                                        ),
                                    ],
                                ),
                                html.Div(
                                    id="title-edit-row",
                                    className="title-edit-row",
                                    children=[
                                        dcc.Input(
                                            id="title-edit-input",
                                            type="text",
                                            placeholder="New title",
                                        ),
                                        html.Button(
                                            "Save",
                                            id="save-title-btn",
                                            className="icon-btn",
                                        ),
                                        html.Button(
                                            "Cancel",
                                            id="cancel-title-btn",
                                            className="icon-btn",
                                        ),
                                    ],
                                ),
                            ],
                        ),
                        html.Div(
                            className="header-actions",
                            children=[
                                dcc.Checklist(
                                    id="theme-toggle",
                                    options=[{"label": "Dark mode", "value": "dark"}],
                                    value=[],
                                    className="theme-toggle",
                                ),
                                html.Div("Jacques Assistant", className="badge"),
                            ],
                        ),
                    ],
                ),
                html.Div(
                    className="content-shell",
                    children=[
                        html.Div(
                            className="chat-stack",
                            children=[
                                html.Div(id="chat-window", className="chat-window"),
                                html.Div(
                                    className="input-panel",
                                    children=[
                                        html.Div(id="tool-status", className="tool-status"),
                                        dcc.Upload(
                                            id="upload-assets",
                                            className="upload-wrap",
                                            multiple=True,
                                            disable_click=True,
                                            accept="image/*,.pdf,.docx,.xlsx,.xls,.xlsm,.csv",
                                            children=html.Div(
                                                className="chat-input upload-zone",
                                                children=[
                                                    html.Div(
                                                        className="chat-textarea-wrap",
                                                        children=[
                                                            dcc.Textarea(
                                                                id="user-input",
                                                                placeholder="Ask Jacques or drop files here",
                                                            ),
                                                            html.Div(
                                                                id="mention-suggestions",
                                                                className="mention-suggestions",
                                                            ),
                                                        ],
                                                    ),
                                                    html.Div(
                                                        className="chat-actions",
                                                        children=[
                                                            html.Button(
                                                                "Send",
                                                                id="send-btn",
                                                                className="send-btn",
                                                            ),
                                                            html.Button(
                                                                "Stop",
                                                                id="stop-btn",
                                                                className="stop-btn",
                                                            ),
                                                        ],
                                                    ),
                                                    html.Div(
                                                        "Drag & drop files here",
                                                        className="upload-hint",
                                                    ),
                                                ],
                                            ),
                                        ),
                                        html.Div(id="upload-status", className="status"),
                                        html.Div(id="stream-status", className="status"),
                                        html.Div(id="action-status", className="status"),
                                        html.Div(id="email-status", className="status"),
                                        html.Div(id="calendar-status", className="status"),
                                    ],
                                ),
                            ],
                        ),
                        html.Div(
                            className="side-panel",
                            children=[
                                html.Div(
                                    id="files-list",
                                    className="files-list",
                                )
                            ],
                        ),
                    ],
                ),
                html.Div(
                    id="file-offcanvas",
                    className="offcanvas",
                    children=[
                        html.Div(
                            className="offcanvas-header",
                            children=[
                                html.Div(
                                    id="file-editor-title",
                                    className="offcanvas-title",
                                ),
                                html.Button(
                                    "Close",
                                    id="close-file-btn",
                                    className="icon-btn",
                                ),
                            ],
                        ),
                        html.Div(
                            id="file-editor-body",
                            className="offcanvas-body",
                        ),
                        html.Div(id="file-save-status", className="status"),
                    ],
                ),
            ],
        ),
        html.Div(
            id="settings-modal",
            className="modal",
            children=[
                html.Div(id="settings-backdrop", className="modal-backdrop"),
                html.Div(
                    className="modal-card",
                    children=[
                        html.Div(
                            className="modal-header",
                            children=[
                                html.Div(
                                    "Settings",
                                    className="modal-title",
                                ),
                                html.Button(
                                    "X",
                                    id="close-settings-btn",
                                    className="icon-btn",
                                ),
                            ],
                        ),
                        html.Div(
                            className="modal-body",
                            children=[
                                dcc.Tabs(
                                    id="settings-tabs",
                                    value="personalization",
                                    className="settings-tabs",
                                    parent_className="settings-tabs-root",
                                    content_className="settings-tabs-content",
                                    vertical=True,
                                    children=[
                                        dcc.Tab(
                                            label="Personalization",
                                            value="personalization",
                                            className="settings-tab",
                                            selected_className="settings-tab--selected",
                                            children=[
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Your profile",
                                                            className="section-title",
                                                        ),
                                                        html.Div(
                                                            className="settings-grid",
                                                            children=[
                                                                dcc.Input(
                                                                    id="user-nickname-input",
                                                                    value=INITIAL_USER_NICKNAME,
                                                                    type="text",
                                                                    placeholder="Nickname",
                                                                    className="settings-input",
                                                                ),
                                                                dcc.Input(
                                                                    id="user-occupation-input",
                                                                    value=INITIAL_USER_OCCUPATION,
                                                                    type="text",
                                                                    placeholder="Occupation",
                                                                    className="settings-input",
                                                                ),
                                                            ],
                                                        ),
                                                        dcc.Textarea(
                                                            id="user-about-input",
                                                            value=INITIAL_USER_ABOUT,
                                                            placeholder="More about you",
                                                            className="settings-textarea",
                                                            rows=3,
                                                        ),
                                                        html.Div(
                                                            "Custom instructions",
                                                            className="field-label",
                                                        ),
                                                        dcc.Textarea(
                                                            id="custom-instructions-input",
                                                            value=INITIAL_CUSTOM_INSTRUCTIONS,
                                                            placeholder="How should Jacques respond?",
                                                            className="settings-textarea settings-textarea--prompt",
                                                            rows=4,
                                                        ),
                                                        html.Button(
                                                            "Save personalization",
                                                            id="save-personalization-btn",
                                                            className="settings-btn",
                                                        ),
                                                        html.Div(
                                                            id="personalization-status",
                                                            className="status",
                                                        ),
                                                    ],
                                                ),
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Advanced tool access",
                                                            className="section-title",
                                                        ),
                                                        dcc.Checklist(
                                                            id="tool-toggles-input",
                                                            options=[
                                                                {
                                                                    "label": "Web search & news",
                                                                    "value": "web",
                                                                },
                                                                {
                                                                    "label": "Mail tools",
                                                                    "value": "mail",
                                                                },
                                                                {
                                                                    "label": "Calendar tools",
                                                                    "value": "calendar",
                                                                },
                                                                {
                                                                    "label": "Code tools",
                                                                    "value": "code",
                                                                },
                                                            ],
                                                            value=[
                                                                option
                                                                for option, enabled in [
                                                                    ("web", TOOLS_WEB_ENABLED),
                                                                    ("mail", TOOLS_MAIL_ENABLED),
                                                                    ("calendar", TOOLS_CALENDAR_ENABLED),
                                                                    ("code", TOOLS_CODE_ENABLED),
                                                                ]
                                                                if enabled
                                                            ],
                                                            className="settings-toggle-list",
                                                        ),
                                                        html.Button(
                                                            "Save tool settings",
                                                            id="save-tool-toggles-btn",
                                                            className="settings-btn secondary",
                                                        ),
                                                        html.Div(
                                                            id="tool-settings-status",
                                                            className="status",
                                                        ),
                                                    ],
                                                ),
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "System prompt",
                                                            className="section-title",
                                                        ),
                                                        dcc.Textarea(
                                                            id="system-prompt-input",
                                                            value=INITIAL_SYSTEM_PROMPT,
                                                            className="settings-textarea settings-textarea--prompt",
                                                            rows=6,
                                                        ),
                                                        html.Button(
                                                            "Save prompt",
                                                            id="save-prompt-btn",
                                                            className="settings-btn secondary",
                                                        ),
                                                        html.Div(
                                                            id="prompt-status",
                                                            className="status",
                                                        ),
                                                        html.Div(
                                                            "Global memory",
                                                            className="field-label",
                                                        ),
                                                        dcc.Textarea(
                                                            id="global-memory-input",
                                                            value=INITIAL_GLOBAL_MEMORY,
                                                            className="settings-textarea settings-textarea--memory",
                                                            rows=5,
                                                        ),
                                                        html.Button(
                                                            "Save memory",
                                                            id="save-memory-btn",
                                                            className="settings-btn",
                                                        ),
                                                        html.Div(
                                                            id="memory-status",
                                                            className="status",
                                                        ),
                                                        html.Div(
                                                            "Saved memory (future chats)",
                                                            className="field-label",
                                                        ),
                                                        html.Div(
                                                            id="memory-preview",
                                                            className="memory-preview",
                                                        ),
                                                    ],
                                                ),
                                            ],
                                        ),
                                        dcc.Tab(
                                            label="App settings",
                                            value="app-settings",
                                            className="settings-tab",
                                            selected_className="settings-tab--selected",
                                            children=[
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Runtime",
                                                            className="section-title",
                                                        ),
                                                        dcc.Checklist(
                                                            id="streaming-toggle-input",
                                                            options=[
                                                                {
                                                                    "label": "LLM streaming",
                                                                    "value": "stream",
                                                                }
                                                            ],
                                                            value=(
                                                                ["stream"]
                                                                if LLM_STREAMING_ENABLED
                                                                else []
                                                            ),
                                                            className="settings-toggle-list",
                                                        ),
                                                        dcc.Input(
                                                            id="app-timezone-input",
                                                            value=INITIAL_APP_TIMEZONE,
                                                            type="text",
                                                            placeholder="Timezone (ex: Europe/Zurich)",
                                                            className="settings-input",
                                                        ),
                                                        html.Button(
                                                            "Save app settings",
                                                            id="save-app-settings-btn",
                                                            className="settings-btn",
                                                        ),
                                                        html.Div(
                                                            id="app-settings-status",
                                                            className="status",
                                                        ),
                                                    ],
                                                ),
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Mail & calendar defaults",
                                                            className="section-title",
                                                        ),
                                                        dcc.Input(
                                                            id="default-mail-account-input",
                                                            value=INITIAL_MAIL_ACCOUNT,
                                                            type="text",
                                                            placeholder="Default mail account (optional)",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="default-mailbox-input",
                                                            value=INITIAL_MAILBOX,
                                                            type="text",
                                                            placeholder="Default mailbox (ex: Inbox)",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="default-calendar-input",
                                                            value=INITIAL_DEFAULT_CALENDAR,
                                                            type="text",
                                                            placeholder="Default calendar name (optional)",
                                                            className="settings-input",
                                                        ),
                                                    ],
                                                ),
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "OnlyOffice",
                                                            className="section-title",
                                                        ),
                                                        dcc.Input(
                                                            id="onlyoffice-url-input",
                                                            value=INITIAL_ONLYOFFICE_URL,
                                                            type="text",
                                                            placeholder="OnlyOffice URL (http://localhost:8080)",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="onlyoffice-jwt-input",
                                                            value=INITIAL_ONLYOFFICE_JWT,
                                                            type="text",
                                                            placeholder="OnlyOffice JWT secret",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="app-base-url-input",
                                                            value=INITIAL_APP_BASE_URL,
                                                            type="text",
                                                            placeholder="Jacques base URL (callback)",
                                                            className="settings-input",
                                                        ),
                                                        html.Button(
                                                            "Save OnlyOffice settings",
                                                            id="save-onlyoffice-btn",
                                                            className="settings-btn secondary",
                                                        ),
                                                        html.Button(
                                                            "Start OnlyOffice (Docker)",
                                                            id="start-onlyoffice-btn",
                                                            className="settings-btn",
                                                        ),
                                                        html.Button(
                                                            "Start Docker Desktop",
                                                            id="start-docker-btn",
                                                            className="settings-btn secondary",
                                                        ),
                                                        html.Div(
                                                            id="onlyoffice-status",
                                                            className="status",
                                                        ),
                                                    ],
                                                ),
                                            ],
                                        ),
                                        dcc.Tab(
                                            label="Tasks",
                                            value="tasks",
                                            className="settings-tab",
                                            selected_className="settings-tab--selected",
                                            children=[
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Scheduled tasks",
                                                            className="section-title",
                                                        ),
                                                        dcc.Dropdown(
                                                            id="task-type-input",
                                                            options=[
                                                                {
                                                                    "label": "Reminder",
                                                                    "value": "reminder",
                                                                },
                                                                {
                                                                    "label": "Web digest",
                                                                    "value": "web_digest",
                                                                },
                                                            ],
                                                            value="reminder",
                                                            clearable=False,
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="task-name-input",
                                                            type="text",
                                                            placeholder="Task name (optional)",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Textarea(
                                                            id="task-message-input",
                                                            placeholder="Reminder message or search query",
                                                            className="settings-textarea",
                                                            rows=3,
                                                        ),
                                                        dcc.Input(
                                                            id="task-cron-input",
                                                            type="text",
                                                            placeholder="Cron: min hour day month dow (ex: 0 8 * * *)",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="task-timezone-input",
                                                            type="text",
                                                            placeholder=f"Timezone (default {settings.app_timezone})",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Input(
                                                            id="task-limit-input",
                                                            type="number",
                                                            placeholder="Web digest count (optional)",
                                                            className="settings-input",
                                                        ),
                                                        dcc.Checklist(
                                                            id="task-use-llm-input",
                                                            options=[
                                                                {
                                                                    "label": "AI summary",
                                                                    "value": "use",
                                                                }
                                                            ],
                                                            value=["use"],
                                                            className="theme-toggle",
                                                        ),
                                                        html.Button(
                                                            "Add task",
                                                            id="task-create-btn",
                                                            className="settings-btn",
                                                        ),
                                                        html.Div(
                                                            id="task-status",
                                                            className="status",
                                                        ),
                                                        html.Div(
                                                            id="tasks-list",
                                                            className="tasks-list",
                                                        ),
                                                    ],
                                                )
                                            ],
                                        ),
                                        dcc.Tab(
                                            label="Data controls",
                                            value="data-controls",
                                            className="settings-tab",
                                            selected_className="settings-tab--selected",
                                            children=[
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Archive & export",
                                                            className="section-title",
                                                        ),
                                                        html.Button(
                                                            "Archive current chat",
                                                            id="archive-current-btn",
                                                            className="settings-btn secondary",
                                                        ),
                                                        html.Button(
                                                            "Archive all chats",
                                                            id="archive-all-btn",
                                                            className="settings-btn secondary",
                                                        ),
                                                        html.Button(
                                                            "Delete all chats",
                                                            id="delete-all-btn",
                                                            className="settings-btn danger",
                                                        ),
                                                        html.Button(
                                                            "Export data",
                                                            id="export-data-btn",
                                                            className="settings-btn",
                                                        ),
                                                        html.Div(
                                                            id="data-status",
                                                            className="status",
                                                        ),
                                                        html.Div(
                                                            id="export-link",
                                                            className="status",
                                                        ),
                                                    ],
                                                )
                                            ],
                                        ),
                                        dcc.Tab(
                                            label="Archived chats",
                                            value="archived-chats",
                                            className="settings-tab",
                                            selected_className="settings-tab--selected",
                                            children=[
                                                html.Div(
                                                    className="settings-section",
                                                    children=[
                                                        html.Div(
                                                            "Archived conversations",
                                                            className="section-title",
                                                        ),
                                                        html.Div(
                                                            id="archived-list",
                                                            className="archived-list",
                                                        ),
                                                    ],
                                                )
                                            ],
                                        ),
                                    ],
                                ),
                            ],
                        ),
                    ],
                ),
            ],
        ),
    ],
)


@app.callback(
    Output("convo-dropdown", "options", allow_duplicate=True),
    Output("convo-dropdown", "value", allow_duplicate=True),
    Input("new-convo-btn", "n_clicks"),
    prevent_initial_call=True,
)
def create_conversation(n_clicks: int):
    if not n_clicks:
        return no_update, no_update
    clean_title = f"Conversation {len(db.list_conversations()) + 1}"
    convo_id = db.create_conversation(clean_title)
    return _conversation_options(), str(convo_id)


@app.callback(
    Output("chat-window", "children"),
    Output("chat-title", "children"),
    Input("convo-dropdown", "value"),
    Input("chat-refresh", "data"),
)
def refresh_chat(convo_id: str | None, refresh: int):
    if not convo_id:
        return [html.Div("Select a conversation.", className="status")], ""
    conversation = db.get_conversation(int(convo_id))
    rows = db.get_messages(int(convo_id))
    title = conversation["title"] if conversation else "Conversation"
    return _render_messages(rows), title


@app.callback(
    Output("convo-dropdown", "options", allow_duplicate=True),
    Input("chat-refresh", "data"),
    prevent_initial_call=True,
)
def refresh_conversation_options(refresh_value: int):
    return _conversation_options()


@app.callback(
    Output("title-edit-row", "className"),
    Output("title-edit-input", "value"),
    Output("chat-refresh", "data", allow_duplicate=True),
    Output("convo-dropdown", "options", allow_duplicate=True),
    Input("edit-title-btn", "n_clicks"),
    Input("cancel-title-btn", "n_clicks"),
    Input("save-title-btn", "n_clicks"),
    Input("convo-dropdown", "value"),
    State("title-edit-input", "value"),
    State("chat-refresh", "data"),
    prevent_initial_call=True,
)
def edit_conversation_title(
    edit_clicks: int | None,
    cancel_clicks: int | None,
    save_clicks: int | None,
    convo_id: str | None,
    new_title: str | None,
    refresh_value: int,
):
    trigger = dash.callback_context.triggered_id
    if trigger == "edit-title-btn":
        if not convo_id:
            return "title-edit-row", "", no_update, no_update
        conversation = db.get_conversation(int(convo_id))
        title = conversation["title"] if conversation else ""
        return "title-edit-row title-edit-row--open", title, no_update, no_update
    if trigger in {"cancel-title-btn", "convo-dropdown"}:
        return "title-edit-row", "", no_update, no_update
    if trigger == "save-title-btn":
        if not convo_id:
            return "title-edit-row", "", no_update, no_update
        title = (new_title or "").strip()
        if not title:
            return "title-edit-row", "", no_update, no_update
        db.update_conversation_title(int(convo_id), title, auto_title=False)
        return (
            "title-edit-row",
            "",
            (refresh_value or 0) + 1,
            _conversation_options(),
        )
    return "title-edit-row", "", no_update, no_update


@app.callback(
    Output("chat-refresh", "data", allow_duplicate=True),
    Output("action-status", "children"),
    Output("user-input", "value"),
    Input("send-btn", "n_clicks"),
    State("chat-refresh", "data"),
    State("convo-dropdown", "value"),
    State("user-input", "value"),
    prevent_initial_call=True,
)
def send_message(
    n_clicks: int,
    refresh_value: int,
    convo_id: str | None,
    user_text: str | None,
):
    if not n_clicks:
        return no_update, no_update, no_update
    message = (user_text or "").strip()
    if not message:
        return no_update, "Type a message first.", ""
    if not convo_id:
        convo_id = str(db.create_conversation("Conversation 1"))

    use_rag = bool(db.list_documents(int(convo_id)))
    use_web = _setting_enabled("tools_web_enabled", True)

    conversation_id = int(convo_id)
    if _is_streaming(conversation_id):
        return no_update, "A response is already streaming.", ""

    db.add_message(conversation_id, "user", message)
    _set_cancelled(conversation_id, False)

    def on_tool_event(name: str, stage: str) -> None:
        _set_tool_status(conversation_id, name, stage)

    if settings.llm_streaming:
        assistant_message_id = db.add_message(conversation_id, "assistant", "")
        _set_streaming(conversation_id, True)
        thread = threading.Thread(
            target=_stream_response,
            args=(
                conversation_id,
                message,
                use_rag,
                use_web,
                assistant_message_id,
            ),
            daemon=True,
        )
        thread.start()
        return refresh_value + 1, "Streaming response...", ""

    reply = assistant.respond(
        conversation_id,
        message,
        settings,
        use_rag=use_rag,
        use_web=use_web,
        on_tool_event=on_tool_event,
        should_cancel=lambda: _is_cancelled(conversation_id),
    )
    db.add_message(conversation_id, "assistant", reply)
    assistant.maybe_update_conversation_title(
        conversation_id, settings, force_first=True
    )
    return refresh_value + 1, "Response generated.", ""


@app.callback(
    Output("action-status", "children", allow_duplicate=True),
    Input("stop-btn", "n_clicks"),
    State("convo-dropdown", "value"),
    prevent_initial_call=True,
)
def stop_response(n_clicks: int | None, convo_id: str | None):
    if not n_clicks:
        return no_update
    if not convo_id:
        return "Select a conversation first."
    conversation_id = int(convo_id)
    if not _is_streaming(conversation_id):
        return "No active response to stop."
    _set_cancelled(conversation_id, True)
    _set_tool_status(conversation_id, None, "idle")
    return "Stopping..."


@app.callback(
    Output("docs-refresh", "data"),
    Output("images-refresh", "data"),
    Output("upload-status", "children"),
    Input("upload-assets", "contents"),
    State("upload-assets", "filename"),
    State("convo-dropdown", "value"),
    State("docs-refresh", "data"),
    State("images-refresh", "data"),
    prevent_initial_call=True,
)
def upload_assets(contents_list, filenames, convo_id, docs_refresh, images_refresh):
    if not contents_list or not filenames:
        return no_update, no_update, no_update
    if not convo_id:
        convo_id = str(db.create_conversation("Conversation 1"))
    conversation_id = int(convo_id)

    doc_updated = False
    image_updated = False
    messages = []
    image_exts = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
    doc_exts = {".pdf", ".docx", ".xlsx", ".xls", ".xlsm", ".csv"}

    for contents, filename in zip(contents_list, filenames):
        safe_name = safe_filename(filename)
        data, _ = decode_upload(contents)
        suffix = Path(safe_name).suffix.lower()
        if suffix in image_exts:
            path = IMAGES_DIR / safe_name
            path.write_bytes(data)
            try:
                description = vision.describe_image(path, settings)
            except Exception as exc:
                description = f"Vision failed: {exc}"
            db.add_image(conversation_id, safe_name, str(path), description)
            image_updated = True
            messages.append(f"Stored {safe_name}")
        elif suffix in doc_exts:
            path = UPLOADS_DIR / safe_name
            path.write_bytes(data)
            try:
                text = doc_ingest.extract_text(path)
                db.add_document(
                    conversation_id, safe_name, str(path), path.suffix.lower(), text
                )
                doc_updated = True
                messages.append(f"Ingested {safe_name}")
            except Exception as exc:
                messages.append(f"Failed {safe_name}: {exc}")
        else:
            messages.append(f"Unsupported file: {safe_name}")

    if doc_updated:
        rag.build_index(conversation_id)

    docs_next = (docs_refresh or 0) + 1 if doc_updated else no_update
    images_next = (images_refresh or 0) + 1 if image_updated else no_update
    status = " | ".join(messages) if messages else "Upload complete."
    return docs_next, images_next, status


@app.callback(
    Output("files-list", "children"),
    Input("docs-refresh", "data"),
    Input("images-refresh", "data"),
    Input("convo-dropdown", "value"),
)
def refresh_files(docs_refresh: int, images_refresh: int, convo_id: str | None):
    if not convo_id:
        return html.Div("No conversation selected.", className="status")
    conversation_id = int(convo_id)
    docs = db.list_documents(conversation_id)
    images = db.list_images(conversation_id)
    if not docs and not images:
        return html.Div("No files yet.", className="status")

    items = []
    for doc in docs:
        label, kind = _doc_label(doc["doc_type"])
        detail = doc["doc_type"].lstrip(".").upper()
        items.append(
            _file_item(
                label,
                doc["name"],
                detail,
                {"type": "file-item", "index": f"doc-{doc['id']}"},
            )
        )

    for img in images:
        items.append(
            _file_item(
                "IMG",
                img["name"],
                "IMAGE",
                {"type": "file-item", "index": f"img-{img['id']}"},
            )
        )

    return html.Div(
        [
            html.Div("Fichiers", className="sources-title"),
            html.Div(items, className="files-items"),
        ],
        className="files-block",
    )


@app.callback(
    Output("docs-list", "data"),
    Input("docs-refresh", "data"),
    Input("convo-dropdown", "value"),
)
def refresh_docs_list(docs_refresh: int, convo_id: str | None):
    return _get_docs_for_conversation(convo_id)


@app.callback(
    Output("active-file", "data"),
    Input({"type": "file-item", "index": ALL}, "n_clicks"),
    Input("close-file-btn", "n_clicks"),
    State("convo-dropdown", "value"),
    State({"type": "file-item", "index": ALL}, "id"),
    prevent_initial_call=True,
)
def select_file(
    file_clicks: list[int],
    close_clicks: int | None,
    convo_id: str | None,
    file_ids: list[dict] | None,
):
    trigger = dash.callback_context.triggered_id
    if trigger == "close-file-btn":
        return None
    if not isinstance(trigger, dict) or trigger.get("type") != "file-item":
        return no_update
    click_value = None
    for item_id, clicks in zip(file_ids or [], file_clicks or []):
        if item_id == trigger:
            click_value = clicks
            break
    if not click_value:
        return no_update
    if not convo_id:
        return no_update

    index = str(trigger.get("index", ""))
    if index.startswith("doc-"):
        doc_id = int(index.split("-", 1)[1])
        doc = db.get_document_by_id(doc_id)
        if not doc or str(doc["conversation_id"]) != str(convo_id):
            return no_update
        return {
            "kind": "doc",
            "id": doc_id,
            "name": doc["name"],
            "path": doc["path"],
            "doc_type": doc["doc_type"],
        }
    if index.startswith("img-"):
        image_id = int(index.split("-", 1)[1])
        image = db.get_image_by_id(image_id)
        if not image or str(image["conversation_id"]) != str(convo_id):
            return no_update
        return {
            "kind": "img",
            "id": image_id,
            "name": image["name"],
            "path": image["path"],
        }
    return no_update


@app.callback(
    Output("file-offcanvas", "className"),
    Output("file-editor-title", "children"),
    Output("file-editor-body", "children"),
    Input("active-file", "data"),
)
def render_file_offcanvas(active_file: dict | None):
    if not active_file:
        return "offcanvas", "", ""

    name = active_file.get("name") or "File"
    kind = active_file.get("kind")
    path_str = active_file.get("path") or ""
    path = Path(path_str)
    if path_str and not path.exists():
        return (
            "offcanvas offcanvas--open",
            name,
            html.Div("File not found on disk.", className="status"),
        )

    if kind == "img":
        url = _data_url(path_str)
        body = html.Div(
            [
                html.Img(src=url, className="file-image"),
                html.Div(name, className="file-meta"),
            ],
            className="file-viewer",
        )
        return "offcanvas offcanvas--open", name, body

    doc_type = (active_file.get("doc_type") or path.suffix).lower()
    if doc_type == ".pdf":
        file_url = _data_url(path_str)
        highlight_url = (
            f"/assets/pdf_viewer.html?file={file_url}&doc_id={active_file.get('id')}&embed=1"
        )
        body = html.Div(
            [
                html.Iframe(
                    src=highlight_url,
                    className="pdf-frame",
                )
            ],
            className="file-viewer",
        )
        return "offcanvas offcanvas--open", name, body

    if doc_type == ".docx":
        if settings.onlyoffice_url:
            iframe_src = f"/onlyoffice/{active_file.get('id')}"
            body = html.Div(
                [
                    html.Div(
                        "OnlyOffice editor actif.",
                        className="status",
                    ),
                    html.Iframe(src=iframe_src, className="office-frame"),
                ],
                className="file-viewer",
            )
            return "offcanvas offcanvas--open", name, body

        try:
            preview_component = _render_word_preview(path)
            status = (
                "OnlyOffice requis pour l'edition fidele. Configurez-le pour modifier ce document."
            )
        except Exception as exc:
            preview_component = html.Pre("", className="file-preview")
            status = f"Word load failed: {exc}"
        body = html.Div(
            [
                html.Div("Preview (read-only)", className="field-label"),
                preview_component,
                html.Div(status, className="status"),
            ],
            className="file-editor",
        )
        return "offcanvas offcanvas--open", name, body

    if doc_type in {".xlsx", ".xls", ".xlsm", ".csv"}:
        if doc_type != ".csv" and settings.onlyoffice_url:
            iframe_src = f"/onlyoffice/{active_file.get('id')}"
            body = html.Div(
                [
                    html.Div(
                        "OnlyOffice editor actif.",
                        className="status",
                    ),
                    html.Iframe(src=iframe_src, className="office-frame"),
                ],
                className="file-viewer",
            )
            return "offcanvas offcanvas--open", name, body

        sheet_value = None
        try:
            df, sheets = _load_table_data(path)
            if sheets:
                sheet_value = sheets[0]
                df, _ = _load_table_data(path, sheet_name=sheet_value)
            status = ""
        except Exception as exc:
            df = None
            sheets = []
            status = f"Table load failed: {exc}"

        if df is None:
            body = html.Div(status, className="status")
        else:
            df = df.where(df.notna(), "")
            columns = [{"name": str(col), "id": str(col)} for col in df.columns]
            doc_id = active_file.get("id")
            table = dash_table.DataTable(
                id={"type": "excel-editor", "doc": doc_id},
                data=df.to_dict("records"),
                columns=columns,
                editable=True,
                page_size=12,
                selected_cells=[],
                style_table={"overflowX": "auto"},
                style_cell={
                    "minWidth": "120px",
                    "width": "160px",
                    "maxWidth": "260px",
                    "whiteSpace": "normal",
                },
            )
            sheet_select = ""
            if sheets:
                sheet_select = dcc.Dropdown(
                    id={"type": "excel-sheet", "doc": doc_id},
                    options=[{"label": sheet, "value": sheet} for sheet in sheets],
                    value=sheet_value,
                    clearable=False,
                    className="sheet-selector",
                )
            body = html.Div(
                [
                    sheet_select,
                    dcc.Store(id={"type": "excel-selection", "doc": doc_id}, data=""),
                    table,
                    html.Div(
                        id={"type": "excel-preview", "doc": doc_id},
                        className="status",
                    ),
                    html.Button(
                        "Envoyer la selection au chat",
                        id={"type": "excel-action", "name": "send-selection", "doc": doc_id},
                        className="icon-btn",
                    ),
                    html.Button(
                        "Save changes",
                        id={"type": "excel-action", "name": "save", "doc": doc_id},
                        className="send-btn",
                    ),
                    html.Div(status, className="status"),
                ],
                className="file-editor",
            )
        return "offcanvas offcanvas--open", name, body

    return "offcanvas offcanvas--open", name, html.Div(
        "Unsupported file type.", className="status"
    )


@app.callback(
    Output({"type": "excel-editor", "doc": MATCH}, "data"),
    Output({"type": "excel-editor", "doc": MATCH}, "columns"),
    Input({"type": "excel-sheet", "doc": MATCH}, "value"),
    State("active-file", "data"),
    prevent_initial_call=True,
)
def switch_excel_sheet(sheet_name: str | None, active_file: dict | None):
    if not active_file or not sheet_name:
        return no_update, no_update
    path = Path(active_file.get("path") or "")
    try:
        df, _ = _load_table_data(path, sheet_name=sheet_name)
    except Exception:
        return no_update, no_update
    df = df.where(df.notna(), "")
    columns = [{"name": str(col), "id": str(col)} for col in df.columns]
    return df.to_dict("records"), columns


@app.callback(
    Output({"type": "excel-selection", "doc": MATCH}, "data"),
    Output({"type": "excel-preview", "doc": MATCH}, "children"),
    Input({"type": "excel-editor", "doc": MATCH}, "selected_cells"),
    State({"type": "excel-editor", "doc": MATCH}, "data"),
    State({"type": "excel-editor", "doc": MATCH}, "columns"),
    prevent_initial_call=True,
)
def update_excel_selection(
    selected_cells: list[dict] | None,
    table_data: list[dict] | None,
    table_columns: list[dict] | None,
):
    if not selected_cells:
        return "", "Aucune selection."
    if len(selected_cells) > 200:
        return "", "Selection trop grande (max 200 cellules)."

    data = table_data or []
    columns = [col["id"] for col in (table_columns or [])]
    lines = []
    for cell in selected_cells:
        row_idx = int(cell.get("row", 0))
        col_id = cell.get("column_id")
        value = ""
        if 0 <= row_idx < len(data):
            value = data[row_idx].get(col_id, "")
        lines.append(f"R{row_idx + 2} {col_id}: {value}")
    text = "Excel selection:\n" + "\n".join(lines)
    preview = f"{len(selected_cells)} cellules selectionnees."
    return text, preview


@app.callback(
    Output("user-input", "value", allow_duplicate=True),
    Input({"type": "excel-action", "name": "send-selection", "doc": ALL}, "n_clicks"),
    State({"type": "excel-selection", "doc": ALL}, "data"),
    State({"type": "excel-selection", "doc": ALL}, "id"),
    State("user-input", "value"),
    prevent_initial_call=True,
)
def send_excel_selection(
    n_clicks: list[int] | None,
    selection_texts: list[str] | None,
    selection_ids: list[dict] | None,
    current_text: str | None,
):
    trigger = dash.callback_context.triggered_id
    if not trigger or not isinstance(trigger, dict):
        return no_update
    doc_id = trigger.get("doc")
    selection_text = ""
    for item_id, text in zip(selection_ids or [], selection_texts or []):
        if item_id.get("doc") == doc_id:
            selection_text = text or ""
            break
    if not selection_text:
        return no_update
    current = current_text or ""
    prefix = "\n\n" if current.strip() else ""
    return f"{current}{prefix}{selection_text}\n"


@app.callback(
    Output("file-save-status", "children", allow_duplicate=True),
    Output("docs-refresh", "data", allow_duplicate=True),
    Input({"type": "word-action", "name": ALL, "doc": ALL}, "n_clicks"),
    State({"type": "word-input", "name": ALL, "doc": ALL}, "value"),
    State({"type": "word-input", "name": ALL, "doc": ALL}, "id"),
    State("active-file", "data"),
    State("docs-refresh", "data"),
    prevent_initial_call=True,
)
def save_word_edits(
    clicks: list[int] | None,
    input_values: list[str] | None,
    input_ids: list[dict] | None,
    active_file: dict | None,
    docs_refresh: int,
):
    trigger = dash.callback_context.triggered_id
    if not active_file or not trigger or not isinstance(trigger, dict):
        return no_update, no_update
    if active_file.get("kind") != "doc":
        return "Only documents can be edited.", no_update

    doc_id = int(active_file.get("id") or 0)
    doc_row = db.get_document_by_id(doc_id)
    if not doc_row:
        return "Document not found.", no_update

    path = Path(doc_row["path"])
    doc_type = (doc_row["doc_type"] or path.suffix).lower()
    if doc_type != ".docx":
        return "This file is not a Word document.", no_update

    values = {}
    for item_id, value in zip(input_ids or [], input_values or []):
        name = item_id.get("name")
        values[name] = value

    try:
        action = trigger.get("name")
        if action == "append":
            text = (values.get("append") or "").strip()
            if not text:
                return "Provide text to append.", no_update
            file_ops.append_paragraph(path, text)
        elif action == "replace":
            old = (values.get("find") or "").strip()
            if not old:
                return "Provide text to find.", no_update
            file_ops.replace_text(path, old, values.get("replace") or "")
        else:
            return no_update, no_update
    except Exception as exc:
        return f"Save failed: {exc}", no_update

    try:
        text = doc_ingest.extract_text(path)
        db.update_document_text(doc_id, text)
        rag.build_index(int(doc_row["conversation_id"]))
    except Exception:
        pass

    return "Saved changes.", (docs_refresh or 0) + 1


@app.callback(
    Output("file-save-status", "children", allow_duplicate=True),
    Output("docs-refresh", "data", allow_duplicate=True),
    Input({"type": "excel-action", "name": "save", "doc": ALL}, "n_clicks"),
    State({"type": "excel-editor", "doc": ALL}, "data"),
    State({"type": "excel-editor", "doc": ALL}, "columns"),
    State({"type": "excel-sheet", "doc": ALL}, "value"),
    State("active-file", "data"),
    State("docs-refresh", "data"),
    prevent_initial_call=True,
)
def save_table_edits(
    save_clicks: list[int] | None,
    table_data_list: list[list[dict]] | None,
    table_columns_list: list[list[dict]] | None,
    sheet_name_list: list[str] | None,
    active_file: dict | None,
    docs_refresh: int,
):
    trigger = dash.callback_context.triggered_id
    if not trigger or not active_file:
        return no_update, no_update
    if active_file.get("kind") != "doc":
        return "Only documents can be edited.", no_update

    doc_id = int(active_file.get("id") or 0)
    doc_row = db.get_document_by_id(doc_id)
    if not doc_row:
        return "Document not found.", no_update

    path = Path(doc_row["path"])
    doc_type = (doc_row["doc_type"] or path.suffix).lower()
    if doc_type not in {".xlsx", ".xls", ".xlsm", ".csv"}:
        return "This file is not a table.", no_update

    table_data = (table_data_list or [None])[0]
    table_columns = (table_columns_list or [None])[0]
    sheet_name = (sheet_name_list or [None])[0]

    try:
        columns = [col["id"] for col in (table_columns or [])]
        if doc_type == ".csv":
            try:
                import pandas as pd
            except ImportError as exc:
                return f"Pandas is required: {exc}", no_update
            df = pd.DataFrame(table_data or [], columns=columns)
            df.to_csv(path, index=False)
        else:
            from openpyxl import load_workbook

            workbook = load_workbook(path)
            sheet = sheet_name or workbook.sheetnames[0]
            if sheet not in workbook.sheetnames:
                return "Sheet not found.", no_update
            ws = workbook[sheet]
            max_cols = len(columns)
            max_rows = max(ws.max_row, (len(table_data or []) + 1))
            for col_idx, header in enumerate(columns, start=1):
                ws.cell(row=1, column=col_idx).value = header
            for row_idx in range(2, max_rows + 1):
                row_data = None
                if table_data and row_idx - 2 < len(table_data):
                    row_data = table_data[row_idx - 2]
                for col_idx in range(1, max_cols + 1):
                    header = columns[col_idx - 1]
                    if row_data is None:
                        value = None
                    else:
                        value = _coerce_cell_value(row_data.get(header))
                    ws.cell(row=row_idx, column=col_idx).value = value
            workbook.save(path)
    except Exception as exc:
        return f"Save failed: {exc}", no_update

    try:
        text = doc_ingest.extract_text(path)
        db.update_document_text(doc_id, text)
        rag.build_index(int(doc_row["conversation_id"]))
    except Exception:
        pass

    return "Saved changes.", (docs_refresh or 0) + 1


@app.callback(
    Output("theme-root", "data-theme"),
    Input("theme-toggle", "value"),
)
def toggle_theme(theme_value: list[str] | None):
    return "dark" if theme_value and "dark" in theme_value else "light"


@app.callback(
    Output("settings-open", "data"),
    Input("open-settings-btn", "n_clicks"),
    Input("close-settings-btn", "n_clicks"),
    Input("settings-backdrop", "n_clicks"),
    State("settings-open", "data"),
    prevent_initial_call=True,
)
def toggle_settings_modal(
    open_clicks: int | None,
    close_clicks: int | None,
    backdrop_clicks: int | None,
    is_open: bool,
):
    trigger = dash.callback_context.triggered_id
    if trigger == "open-settings-btn":
        return True
    if trigger in {"close-settings-btn", "settings-backdrop"}:
        return False
    return is_open


@app.callback(
    Output("settings-modal", "className"),
    Input("settings-open", "data"),
)
def render_settings_modal(is_open: bool):
    return "modal modal--open" if is_open else "modal"


@app.callback(
    Output("prompt-status", "children"),
    Output("system-prompt-input", "value"),
    Input("save-prompt-btn", "n_clicks"),
    State("system-prompt-input", "value"),
    prevent_initial_call=True,
)
def save_system_prompt(n_clicks: int, prompt_text: str | None):
    if not n_clicks:
        return no_update, no_update
    prompt = (prompt_text or "").strip()
    if not prompt:
        db.set_setting("system_prompt", "")
        return "System prompt reset to default.", assistant.SYSTEM_PROMPT
    db.set_setting("system_prompt", prompt)
    return "System prompt saved.", prompt


@app.callback(
    Output("memory-status", "children"),
    Output("global-memory-input", "value"),
    Input("save-memory-btn", "n_clicks"),
    State("global-memory-input", "value"),
    prevent_initial_call=True,
)
def save_global_memory(n_clicks: int, memory_text: str | None):
    if not n_clicks:
        return no_update, no_update
    memory = (memory_text or "").strip()
    db.set_setting("global_memory", memory)
    status = "Global memory saved." if memory else "Global memory cleared."
    return status, memory


@app.callback(
    Output("personalization-status", "children"),
    Output("custom-instructions-input", "value"),
    Output("user-nickname-input", "value"),
    Output("user-occupation-input", "value"),
    Output("user-about-input", "value"),
    Input("save-personalization-btn", "n_clicks"),
    State("custom-instructions-input", "value"),
    State("user-nickname-input", "value"),
    State("user-occupation-input", "value"),
    State("user-about-input", "value"),
    prevent_initial_call=True,
)
def save_personalization(
    n_clicks: int | None,
    instructions: str | None,
    nickname: str | None,
    occupation: str | None,
    about: str | None,
):
    if not n_clicks:
        return no_update, no_update, no_update, no_update, no_update
    clean_instructions = (instructions or "").strip()
    clean_nickname = (nickname or "").strip()
    clean_occupation = (occupation or "").strip()
    clean_about = (about or "").strip()
    db.set_setting("custom_instructions", clean_instructions)
    db.set_setting("user_nickname", clean_nickname)
    db.set_setting("user_occupation", clean_occupation)
    db.set_setting("user_about", clean_about)
    status = "Personalization saved."
    if not any([clean_instructions, clean_nickname, clean_occupation, clean_about]):
        status = "Personalization cleared."
    return (
        status,
        clean_instructions,
        clean_nickname,
        clean_occupation,
        clean_about,
    )


@app.callback(
    Output("tool-settings-status", "children"),
    Input("save-tool-toggles-btn", "n_clicks"),
    State("tool-toggles-input", "value"),
    prevent_initial_call=True,
)
def save_tool_settings(n_clicks: int | None, values: list[str] | None):
    if not n_clicks:
        return no_update
    enabled = set(values or [])
    db.set_setting("tools_web_enabled", "true" if "web" in enabled else "false")
    db.set_setting("tools_mail_enabled", "true" if "mail" in enabled else "false")
    db.set_setting(
        "tools_calendar_enabled", "true" if "calendar" in enabled else "false"
    )
    db.set_setting("tools_code_enabled", "true" if "code" in enabled else "false")
    return "Tool access updated."


@app.callback(
    Output("app-settings-status", "children"),
    Output("app-timezone-input", "value"),
    Input("save-app-settings-btn", "n_clicks"),
    State("app-timezone-input", "value"),
    State("streaming-toggle-input", "value"),
    State("default-mail-account-input", "value"),
    State("default-mailbox-input", "value"),
    State("default-calendar-input", "value"),
    prevent_initial_call=True,
)
def save_app_settings(
    n_clicks: int | None,
    timezone_value: str | None,
    streaming_values: list[str] | None,
    default_mail_account: str | None,
    default_mailbox: str | None,
    default_calendar: str | None,
):
    if not n_clicks:
        return no_update, no_update
    timezone = (timezone_value or settings.app_timezone or "UTC").strip()
    streaming_enabled = bool(streaming_values and "stream" in streaming_values)
    mail_account = (default_mail_account or "").strip()
    mail_box = (default_mailbox or "").strip()
    calendar_name = (default_calendar or "").strip()
    db.set_setting("app_timezone", timezone)
    db.set_setting("llm_streaming", "true" if streaming_enabled else "false")
    db.set_setting("default_mail_account", mail_account)
    db.set_setting("default_mailbox", mail_box)
    db.set_setting("default_calendar", calendar_name)
    settings.app_timezone = timezone
    settings.llm_streaming = streaming_enabled
    return "App settings saved.", timezone


@app.callback(
    Output("onlyoffice-status", "children"),
    Output("onlyoffice-url-input", "value"),
    Output("onlyoffice-jwt-input", "value"),
    Output("app-base-url-input", "value"),
    Input("save-onlyoffice-btn", "n_clicks"),
    Input("start-onlyoffice-btn", "n_clicks"),
    Input("start-docker-btn", "n_clicks"),
    State("onlyoffice-url-input", "value"),
    State("onlyoffice-jwt-input", "value"),
    State("app-base-url-input", "value"),
    prevent_initial_call=True,
)
def save_onlyoffice_settings(
    save_clicks: int | None,
    start_clicks: int | None,
    docker_clicks: int | None,
    url_value: str | None,
    jwt_value: str | None,
    base_url_value: str | None,
):
    trigger = dash.callback_context.triggered_id
    url = (url_value or "").strip()
    jwt_secret = (jwt_value or "").strip()
    base_url = (base_url_value or "").strip()
    status = ""

    if trigger == "start-docker-btn":
        ok, message = _start_docker_desktop()
        return message, url, jwt_secret, base_url
    if trigger == "start-onlyoffice-btn":
        if not url:
            url = "http://localhost:8080"
        if not jwt_secret:
            jwt_secret = secrets.token_urlsafe(24)
        if not base_url:
            base_url = _default_docker_base_url()
        if (
            ("localhost" in url or "127.0.0.1" in url)
            and ("localhost" in base_url or "127.0.0.1" in base_url)
        ):
            base_url = _default_docker_base_url()
        ok, message = _start_onlyoffice_container(jwt_secret)
        status = message if ok else f"OnlyOffice start failed: {message}"
    elif trigger == "save-onlyoffice-btn":
        status = "OnlyOffice settings saved." if url else "OnlyOffice disabled."
    else:
        return no_update, no_update, no_update, no_update

    db.set_setting("onlyoffice_url", url)
    db.set_setting("onlyoffice_jwt", jwt_secret)
    db.set_setting("app_base_url", base_url)
    settings.onlyoffice_url = url
    settings.onlyoffice_jwt = jwt_secret
    settings.app_base_url = base_url

    return status, url, jwt_secret, base_url


@app.callback(
    Output("tasks-refresh", "data", allow_duplicate=True),
    Output("task-status", "children", allow_duplicate=True),
    Input("task-create-btn", "n_clicks"),
    State("task-type-input", "value"),
    State("task-name-input", "value"),
    State("task-message-input", "value"),
    State("task-cron-input", "value"),
    State("task-timezone-input", "value"),
    State("task-limit-input", "value"),
    State("task-use-llm-input", "value"),
    State("tasks-refresh", "data"),
    State("convo-dropdown", "value"),
    prevent_initial_call=True,
)
def create_task(
    n_clicks: int | None,
    task_type: str | None,
    name: str | None,
    message: str | None,
    cron: str | None,
    timezone: str | None,
    limit: int | None,
    use_llm_value: list[str] | None,
    refresh_value: int,
    convo_id: str | None,
):
    if not n_clicks:
        return no_update, no_update
    if not convo_id:
        return no_update, "Select a conversation first."
    task_type = (task_type or "reminder").strip()
    cron = (cron or "").strip()
    if not cron:
        return no_update, "Provide a cron schedule (min hour day month dow)."
    timezone = (timezone or settings.app_timezone or "UTC").strip()
    payload: dict[str, str | int | bool] = {}
    clean_name = (name or "").strip()
    use_llm = bool(use_llm_value and "use" in use_llm_value)

    if task_type == "web_digest":
        query = (message or "").strip()
        if not query and not clean_name:
            return no_update, "Provide a search query."
        if not clean_name:
            clean_name = f"Digest: {query}"
        payload = {
            "query": query or clean_name,
            "limit": int(limit) if isinstance(limit, int) and limit > 0 else 5,
            "use_llm": use_llm,
        }
    else:
        text = (message or "").strip()
        if not text and clean_name:
            text = clean_name
        if not text:
            return no_update, "Provide a reminder message."
        if not clean_name:
            clean_name = f"Reminder: {text}"
        payload = {"message": text}
        task_type = "reminder"

    task_id = db.add_scheduled_task(
        conversation_id=int(convo_id),
        name=clean_name,
        task_type=task_type,
        payload=json.dumps(payload, ensure_ascii=True),
        cron=cron,
        timezone=timezone,
        enabled=True,
    )
    try:
        task_scheduler.schedule_task_by_id(task_id, settings)
    except Exception as exc:
        return (refresh_value or 0) + 1, f"Task saved, schedule error: {exc}"
    return (refresh_value or 0) + 1, f"Task created (id {task_id})."


@app.callback(
    Output("tasks-list", "children"),
    Input("tasks-refresh", "data"),
    Input("convo-dropdown", "value"),
)
def refresh_tasks_list(refresh_value: int, convo_id: str | None):
    if not convo_id:
        return html.Div("Select a conversation.", className="status")
    tasks = db.list_scheduled_tasks(int(convo_id))
    if not tasks:
        return html.Div("No scheduled tasks.", className="status")
    items = []
    for task in tasks:
        enabled = bool(task["enabled"])
        summary = _task_payload_summary(task)
        last_run = task["last_run"] or "-"
        last_status = task["last_status"] or "-"
        items.append(
            html.Div(
                [
                    html.Div(
                        [
                            html.Div(task["name"], className="task-name"),
                            html.Div(
                                f"{task['task_type']} · {task['cron']} · {task['timezone']}",
                                className="task-meta",
                            ),
                            html.Div(
                                summary or "No details.",
                                className="task-summary",
                            ),
                            html.Div(
                                f"Last run: {last_run} · {last_status}",
                                className="task-meta",
                            ),
                        ],
                        className="task-info",
                    ),
                    html.Div(
                        [
                            html.Button(
                                "Disable" if enabled else "Enable",
                                id={"type": "task-toggle", "index": task["id"]},
                                className="icon-btn",
                            ),
                            html.Button(
                                "Delete",
                                id={"type": "task-delete", "index": task["id"]},
                                className="icon-btn danger",
                            ),
                        ],
                        className="task-actions",
                    ),
                ],
                className="task-item",
            )
        )
    return items


@app.callback(
    Output("tasks-refresh", "data", allow_duplicate=True),
    Output("task-status", "children", allow_duplicate=True),
    Input({"type": "task-delete", "index": ALL}, "n_clicks"),
    Input({"type": "task-toggle", "index": ALL}, "n_clicks"),
    State("tasks-refresh", "data"),
    State("convo-dropdown", "value"),
    prevent_initial_call=True,
)
def update_task_state(
    delete_clicks: list[int] | None,
    toggle_clicks: list[int] | None,
    refresh_value: int,
    convo_id: str | None,
):
    trigger = dash.callback_context.triggered_id
    if not trigger or not isinstance(trigger, dict):
        return no_update, no_update
    task_id = int(trigger.get("index") or 0)
    if not task_id:
        return no_update, "Task not found."
    task = db.get_scheduled_task(task_id)
    if not task or not convo_id or int(task["conversation_id"]) != int(convo_id):
        return no_update, "Task not found."

    if trigger.get("type") == "task-delete":
        db.delete_scheduled_task(task_id)
        task_scheduler.remove_task(task_id)
        return (refresh_value or 0) + 1, f"Task {task_id} deleted."

    enabled = bool(task["enabled"])
    new_state = not enabled
    db.set_scheduled_task_enabled(task_id, new_state)
    if new_state:
        task_scheduler.schedule_task_by_id(task_id, settings)
    else:
        task_scheduler.remove_task(task_id)
    label = "enabled" if new_state else "disabled"
    return (refresh_value or 0) + 1, f"Task {task_id} {label}."


@app.callback(
    Output("memory-preview", "children"),
    Input("chat-refresh", "data"),
    Input("memory-status", "children"),
)
def refresh_memory_preview(refresh_value: int, memory_status: str | None):
    memory = db.get_setting("global_memory") or ""
    if not memory.strip():
        return html.Div("No saved memory yet.", className="status")
    return dcc.Markdown(memory, link_target="_blank")


@app.callback(
    Output("convo-dropdown", "options", allow_duplicate=True),
    Output("convo-dropdown", "value", allow_duplicate=True),
    Output("convo-status", "children"),
    Input({"type": "delete-convo", "index": ALL}, "n_clicks"),
    State("convo-dropdown", "value"),
    prevent_initial_call=True,
)
def delete_conversation(delete_clicks: list[int], convo_id: str | None):
    trigger = dash.callback_context.triggered_id
    if not trigger or not isinstance(trigger, dict):
        return no_update, no_update, no_update
    conversation_id = int(trigger.get("index", 0) or 0)
    if not conversation_id:
        return no_update, no_update, no_update
    if _is_streaming(conversation_id):
        return no_update, no_update, "Wait for the streaming response to finish."
    _purge_conversation(conversation_id)

    conversations = db.list_conversations()
    if not conversations:
        next_value = str(db.create_conversation("Conversation 1"))
    elif convo_id and str(convo_id) != str(conversation_id):
        next_value = str(convo_id)
    else:
        next_value = str(conversations[0]["id"])

    return _conversation_options(), next_value, "Conversation deleted."


@app.callback(
    Output("convo-dropdown", "options", allow_duplicate=True),
    Output("convo-dropdown", "value", allow_duplicate=True),
    Output("data-status", "children"),
    Output("export-link", "children"),
    Output("archived-refresh", "data", allow_duplicate=True),
    Input("archive-current-btn", "n_clicks"),
    Input("archive-all-btn", "n_clicks"),
    Input("delete-all-btn", "n_clicks"),
    Input("export-data-btn", "n_clicks"),
    State("convo-dropdown", "value"),
    State("archived-refresh", "data"),
    prevent_initial_call=True,
)
def handle_data_controls(
    archive_current: int | None,
    archive_all: int | None,
    delete_all: int | None,
    export_clicks: int | None,
    convo_id: str | None,
    archived_refresh: int,
):
    trigger = dash.callback_context.triggered_id
    if trigger == "export-data-btn":
        export_path = _export_all_data()
        link = html.A(
            export_path.name,
            href=_data_url(str(export_path)),
            target="_blank",
            className="export-link",
        )
        return no_update, no_update, "Export ready.", link, archived_refresh

    if trigger == "archive-current-btn":
        if not convo_id:
            return no_update, no_update, "Select a conversation first.", no_update, archived_refresh
        conversation_id = int(convo_id)
        if _is_streaming(conversation_id):
            return no_update, no_update, "Wait for the streaming response to finish.", no_update, archived_refresh
        db.set_conversation_archived(conversation_id, True)
        conversations = db.list_conversations()
        if not conversations:
            next_id = db.create_conversation("Conversation 1")
            next_value = str(next_id)
        else:
            next_value = str(conversations[0]["id"])
        return (
            _conversation_options(),
            next_value,
            "Conversation archived.",
            no_update,
            (archived_refresh or 0) + 1,
        )

    if trigger == "archive-all-btn":
        active_convos = db.list_conversations(include_archived=True)
        for row in active_convos:
            if _is_streaming(int(row["id"])):
                return (
                    no_update,
                    no_update,
                    "Wait for the streaming response to finish.",
                    no_update,
                    archived_refresh,
                )
        db.archive_all_conversations()
        next_id = db.create_conversation("Conversation 1")
        return (
            _conversation_options(),
            str(next_id),
            "All conversations archived.",
            no_update,
            (archived_refresh or 0) + 1,
        )

    if trigger == "delete-all-btn":
        all_convos = db.list_conversations(include_archived=True)
        for row in all_convos:
            conversation_id = int(row["id"])
            if _is_streaming(conversation_id):
                return (
                    no_update,
                    no_update,
                    "Wait for the streaming response to finish.",
                    no_update,
                    archived_refresh,
                )
        for row in all_convos:
            _purge_conversation(int(row["id"]))
        next_id = db.create_conversation("Conversation 1")
        return (
            _conversation_options(),
            str(next_id),
            "All conversations deleted.",
            no_update,
            (archived_refresh or 0) + 1,
        )

    return no_update, no_update, no_update, no_update, archived_refresh


@app.callback(
    Output("archived-list", "children"),
    Input("archived-refresh", "data"),
    Input("settings-open", "data"),
)
def refresh_archived_list(refresh_value: int, is_open: bool):
    if not is_open:
        return no_update
    archived = db.list_archived_conversations()
    if not archived:
        return html.Div("No archived conversations.", className="status")
    items = []
    for convo in archived:
        items.append(
            html.Div(
                [
                    html.Div(
                        [
                            html.Div(convo["title"], className="archived-title"),
                            html.Div(convo["created_at"], className="archived-meta"),
                        ],
                        className="archived-info",
                    ),
                    html.Div(
                        [
                            html.Button(
                                "Restore",
                                id={"type": "archived-restore", "index": convo["id"]},
                                className="icon-btn",
                            ),
                            html.Button(
                                "Delete",
                                id={"type": "archived-delete", "index": convo["id"]},
                                className="icon-btn danger",
                            ),
                        ],
                        className="archived-actions",
                    ),
                ],
                className="archived-item",
            )
        )
    return items


@app.callback(
    Output("archived-refresh", "data", allow_duplicate=True),
    Output("convo-dropdown", "options", allow_duplicate=True),
    Output("convo-dropdown", "value", allow_duplicate=True),
    Output("data-status", "children", allow_duplicate=True),
    Input({"type": "archived-restore", "index": ALL}, "n_clicks"),
    Input({"type": "archived-delete", "index": ALL}, "n_clicks"),
    State("archived-refresh", "data"),
    State("convo-dropdown", "value"),
    prevent_initial_call=True,
)
def update_archived_items(
    restore_clicks: list[int] | None,
    delete_clicks: list[int] | None,
    refresh_value: int,
    convo_id: str | None,
):
    trigger = dash.callback_context.triggered_id
    if not trigger or not isinstance(trigger, dict):
        return no_update, no_update, no_update, no_update
    convo_id_int = int(trigger.get("index") or 0)
    if not convo_id_int:
        return no_update, no_update, no_update, "Conversation not found."
    if trigger.get("type") == "archived-restore":
        db.set_conversation_archived(convo_id_int, False)
        conversations = db.list_conversations()
        next_value = str(conversations[0]["id"]) if conversations else str(convo_id_int)
        return (
            (refresh_value or 0) + 1,
            _conversation_options(),
            next_value,
            "Conversation restored.",
        )
    if trigger.get("type") == "archived-delete":
        if _is_streaming(convo_id_int):
            return no_update, no_update, no_update, "Wait for the streaming response to finish."
        _purge_conversation(convo_id_int)
        conversations = db.list_conversations()
        next_value = str(conversations[0]["id"]) if conversations else None
        if not next_value:
            next_value = str(db.create_conversation("Conversation 1"))
        return (
            (refresh_value or 0) + 1,
            _conversation_options(),
            next_value,
            "Archived conversation deleted.",
        )
    return no_update, no_update, no_update, no_update


@app.callback(
    Output("tool-status", "children"),
    Input("chat-refresh", "data"),
    Input("stream-interval", "n_intervals"),
    State("convo-dropdown", "value"),
)
def refresh_tool_status(refresh_value: int, n_intervals: int, convo_id: str | None):
    if not convo_id:
        return ""
    conversation_id = int(convo_id)
    state = _get_tool_status(conversation_id) or {}
    tool_name = state.get("name") or _latest_tool_name(conversation_id)
    stage = state.get("stage", "result") if tool_name else "idle"

    if tool_name:
        if stage == "call":
            text = f"Outil en cours: {tool_name}"
            loader_class = "tool-loader"
        else:
            text = f"Dernier outil: {tool_name}"
            loader_class = "tool-loader tool-loader--idle"
    else:
        text = "Aucun outil appele."
        loader_class = "tool-loader tool-loader--idle"
    return html.Div(
        [
            html.Span(text, className="tool-status-text"),
            html.Span(className=loader_class),
        ],
        className="tool-status-row",
    )


@app.callback(
    Output("mention-suggestions", "children"),
    Output("mention-suggestions", "className"),
    Input("user-input", "value"),
    State("docs-list", "data"),
)
def update_mention_suggestions(user_text: str | None, docs_list: list[dict] | None):
    text = user_text or ""
    docs = docs_list or []
    match_text = text.rstrip()
    if not match_text:
        return "", "mention-suggestions"
    match = re.search(r"@([^\\s@]*)$", match_text)
    if not match or not docs:
        return "", "mention-suggestions"
    query = match.group(1).lower()
    candidates = []
    for doc in docs:
        name = str(doc.get("name") or "")
        if not query or query in name.lower():
            candidates.append(doc)
    if not candidates:
        return "", "mention-suggestions"
    buttons = []
    for doc in candidates[:8]:
        name = str(doc.get("name") or "")
        doc_type = str(doc.get("doc_type") or "").lstrip(".").upper()
        buttons.append(
            html.Button(
                [
                    html.Span(name, className="mention-name"),
                    html.Span(doc_type or "DOC", className="mention-tag"),
                ],
                id={"type": "doc-suggestion", "index": doc.get("id")},
                className="mention-item",
                n_clicks=0,
            )
        )
    return buttons, "mention-suggestions mention-suggestions--open"


@app.callback(
    Output("user-input", "value", allow_duplicate=True),
    Input({"type": "doc-suggestion", "index": ALL}, "n_clicks"),
    State("user-input", "value"),
    State("docs-list", "data"),
    prevent_initial_call=True,
)
def apply_doc_suggestion(
    clicks: list[int] | None,
    current_text: str | None,
    docs_list: list[dict] | None,
):
    trigger = dash.callback_context.triggered_id
    if not trigger or not isinstance(trigger, dict):
        return no_update
    doc_id = trigger.get("index")
    docs = docs_list or []
    doc_name = ""
    for doc in docs:
        if doc.get("id") == doc_id:
            doc_name = str(doc.get("name") or "")
            break
    if not doc_name:
        return no_update
    return _apply_doc_mention(current_text or "", doc_name)


@app.callback(
    Output("mailto-data", "data"),
    Input("chat-refresh", "data"),
    Input("stream-interval", "n_intervals"),
    Input("convo-dropdown", "value"),
)
def refresh_mailto_data(
    refresh_value: int, n_intervals: int, convo_id: str | None
):
    if not convo_id:
        return {}
    conversation_id = int(convo_id)
    row = db.get_latest_tool_call_message_by_name(conversation_id, "email_draft")
    if not row:
        return {}
    if not _is_recent_message(row["created_at"], MAILTO_RECENT_SECONDS):
        _should_open_mailto(conversation_id, int(row["id"]))
        return {}
    args = _extract_tool_call_args(row["content"] or "")
    mailto = _build_mailto(args)
    if not mailto:
        return {}
    if not _should_open_mailto(conversation_id, int(row["id"])):
        return {}
    return {
        "convo_id": conversation_id,
        "message_id": int(row["id"]),
        "mailto": mailto,
    }


@app.callback(
    Output("email-status", "children"),
    Input("chat-refresh", "data"),
    Input("convo-dropdown", "value"),
)
def render_email_status(refresh_value: int, convo_id: str | None):
    if not convo_id:
        return ""
    conversation_id = int(convo_id)
    row_call = db.get_latest_tool_call_message_by_name(conversation_id, "email_draft")
    if not row_call:
        return ""
    args = _extract_tool_call_args(row_call["content"] or "")
    mailto = _build_mailto(args)
    if not mailto:
        return ""
    status_line = "Email draft ready."
    row_result = db.get_latest_tool_message_by_name(conversation_id, "email_draft")
    if row_result:
        payload = _extract_tool_result_text(row_result["content"] or "")
        if payload:
            status_line = payload.splitlines()[0]
    to_list = _normalize_list(args.get("to") if isinstance(args, dict) else None)
    subject = ""
    if isinstance(args, dict):
        subject = str(args.get("subject") or "").strip()
    summary = status_line
    if to_list:
        summary += f" · To: {', '.join(to_list)}"
    if subject:
        summary += f" · Subject: {subject}"
    return html.Div(
        [
            html.Span(summary, className="email-status-text"),
            html.A(
                "Open in Mail",
                href=mailto,
                className="icon-btn",
            ),
        ],
        className="email-status-row",
    )


@app.callback(
    Output("calendar-data", "data"),
    Input("chat-refresh", "data"),
    Input("stream-interval", "n_intervals"),
    Input("convo-dropdown", "value"),
)
def refresh_calendar_data(
    refresh_value: int, n_intervals: int, convo_id: str | None
):
    if not convo_id:
        return {}
    conversation_id = int(convo_id)
    row = db.get_latest_tool_message_by_name(conversation_id, "calendar_event")
    if not row:
        return {}
    if not _is_recent_message(row["created_at"], CALENDAR_RECENT_SECONDS):
        _should_open_calendar(conversation_id, int(row["id"]))
        return {}
    ics_link = _extract_ics_link(row["content"] or "")
    if not ics_link:
        return {}
    if not _should_open_calendar(conversation_id, int(row["id"])):
        return {}
    return {
        "convo_id": conversation_id,
        "message_id": int(row["id"]),
        "ics": ics_link,
    }


@app.callback(
    Output("calendar-status", "children"),
    Input("chat-refresh", "data"),
    Input("convo-dropdown", "value"),
)
def render_calendar_status(refresh_value: int, convo_id: str | None):
    if not convo_id:
        return ""
    conversation_id = int(convo_id)
    row = db.get_latest_tool_message_by_name(conversation_id, "calendar_event")
    if not row:
        return ""
    payload = _extract_tool_result_text(row["content"] or "")
    status_line = payload.splitlines()[0] if payload else "Calendar event ready."
    ics_link = _extract_ics_link(row["content"] or "")
    if not ics_link:
        return ""
    return html.Div(
        [
            html.Span(status_line, className="calendar-status-text"),
            html.A(
                "Open in Calendar",
                href=ics_link,
                className="icon-btn",
            ),
        ],
        className="calendar-status-row",
    )


@app.callback(
    Output("chat-refresh", "data", allow_duplicate=True),
    Output("stream-status", "children"),
    Output("stream-state", "data"),
    Input("stream-interval", "n_intervals"),
    State("chat-refresh", "data"),
    State("convo-dropdown", "value"),
    State("stream-state", "data"),
    prevent_initial_call=True,
)
def poll_stream(
    n_intervals: int,
    refresh_value: int,
    convo_id: str | None,
    stream_state: dict | None,
):
    if not convo_id:
        return no_update, "", stream_state or {}
    conversation_id = int(convo_id)
    is_streaming = _is_streaming(conversation_id)
    was_streaming = False
    if stream_state and stream_state.get("convo_id") == convo_id:
        was_streaming = bool(stream_state.get("is_streaming"))

    next_state = {"convo_id": convo_id, "is_streaming": is_streaming}

    if is_streaming:
        return (refresh_value or 0) + 1, "Streaming...", next_state
    if was_streaming and not is_streaming:
        return (refresh_value or 0) + 1, "", next_state
    return no_update, "", next_state


app.clientside_callback(
    """
    function(mailtoData, mailtoState) {
        if (!mailtoData || !mailtoData.mailto) {
            return mailtoState || {};
        }
        const convoId = String(mailtoData.convo_id || "");
        if (!convoId) {
            return mailtoState || {};
        }
        const nextState = Object.assign({}, mailtoState || {});
        if (String(nextState[convoId]) === String(mailtoData.message_id)) {
            return nextState;
        }
        window.location.href = mailtoData.mailto;
        nextState[convoId] = mailtoData.message_id;
        return nextState;
    }
    """,
    Output("mailto-state", "data"),
    Input("mailto-data", "data"),
    State("mailto-state", "data"),
)


app.clientside_callback(
    """
    function(calendarData, calendarState) {
        if (!calendarData || !calendarData.ics) {
            return calendarState || {};
        }
        const convoId = String(calendarData.convo_id || "");
        if (!convoId) {
            return calendarState || {};
        }
        const nextState = Object.assign({}, calendarState || {});
        if (String(nextState[convoId]) === String(calendarData.message_id)) {
            return nextState;
        }
        window.location.href = calendarData.ics;
        nextState[convoId] = calendarData.message_id;
        return nextState;
    }
    """,
    Output("calendar-state", "data"),
    Input("calendar-data", "data"),
    State("calendar-state", "data"),
)


app.clientside_callback(
    """
    function(refreshValue, convoId) {
        const el = document.getElementById("chat-window");
        if (el) {
            el.scrollTop = el.scrollHeight;
        }
        return Date.now();
    }
    """,
    Output("scroll-trigger", "data"),
    Input("chat-refresh", "data"),
    Input("convo-dropdown", "value"),
)


if __name__ == "__main__":
    app.run_server(debug=True)
