from pathlib import Path


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix in {".xlsx", ".xls", ".xlsm"}:
        return _extract_excel(file_path)
    if suffix == ".csv":
        return _extract_csv(file_path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(file_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required for PDF ingestion") from exc

    reader = PdfReader(str(file_path))
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n".join(chunks).strip()


def _extract_docx(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word ingestion") from exc

    doc = Document(str(file_path))
    chunks: list[str] = []

    def add(text: str) -> None:
        cleaned = text.strip()
        if not cleaned:
            return
        if chunks and chunks[-1] == cleaned:
            return
        chunks.append(cleaned)

    for para in doc.paragraphs:
        if para.text:
            add(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    add(cell.text)

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text:
                add(para.text)
        for para in section.footer.paragraphs:
            if para.text:
                add(para.text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        add(cell.text)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        add(cell.text)

    if not chunks:
        try:
            import mammoth
        except ImportError:
            mammoth = None
        if mammoth:
            raw = mammoth.extract_raw_text(str(file_path)).value
            if raw:
                add(raw)

    return "\n".join(chunks).strip()


def _extract_excel(file_path: Path) -> str:
    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError("pandas is required for Excel ingestion") from exc

    sheet_texts = []
    xls = pd.ExcelFile(file_path)
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        sheet_texts.append(f"Sheet: {sheet}")
        sheet_texts.append(df.to_csv(index=False))
    return "\n".join(sheet_texts).strip()


def _extract_csv(file_path: Path) -> str:
    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError("pandas is required for CSV ingestion") from exc

    df = pd.read_csv(file_path)
    return df.to_csv(index=False).strip()
