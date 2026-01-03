from pathlib import Path


def create_excel(path: Path, sheet_name: str = "Sheet1") -> Path:
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for Excel operations") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)
    return path


def add_sheet(path: Path, sheet_name: str) -> Path:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for Excel operations") from exc

    workbook = load_workbook(path)
    workbook.create_sheet(sheet_name)
    workbook.save(path)
    return path


def set_cell(path: Path, sheet_name: str, cell: str, value: str) -> Path:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for Excel operations") from exc

    workbook = load_workbook(path)
    sheet = workbook[sheet_name]
    sheet[cell] = value
    workbook.save(path)
    return path


def create_word(path: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word operations") from exc

    doc = Document()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def append_paragraph(path: Path, text: str) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word operations") from exc

    doc = Document(str(path))
    doc.add_paragraph(text)
    doc.save(path)
    return path


def replace_text(path: Path, old: str, new: str) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word operations") from exc

    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
    doc.save(path)
    return path


def write_word(path: Path, text: str) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word operations") from exc

    doc = Document()
    for line in (text or "").splitlines():
        doc.add_paragraph(line)
    doc.save(path)
    return path
